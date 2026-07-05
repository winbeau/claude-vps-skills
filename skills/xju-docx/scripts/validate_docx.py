# -*- coding: utf-8 -*-
"""validate_docx — 校验/修复 docx 是否符合《新疆大学本科毕业论文（设计）规范及格式要求》。

用法：python3 validate_docx.py <文件.docx> [--framework] [--json] [--fix] [--content]
  --framework  额外检查三节分页结构（封面无页码 → 前置 upperRoman → 正文 decimal 重新起页）
  --json       结构化 JSON 输出（fails 超 20 条截断），供编排脚本/agent 消费
  --fix        先自动修复机械可修项（样式字号字体/页面设置/updateFields；加 --framework 连页码分节一起修），
               回写文件后再机检；不可自动修项列入"需人工确认"
  --content    占位符残留门禁（【待填】/待填/lorem/ipsum/xxx），定稿前必跑

退出码：0 全部通过；1 有 FAIL。
"""
import json
import os
import sys
import zipfile

import docx
from lxml import etree
from docx.oxml.ns import qn

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import xju_format  # noqa: E402

RESULTS = []          # {rule_id, name, ok, detail}
MAX_JSON_FAILS = 20


def check(rule_id, name, ok, detail=''):
    RESULTS.append({'rule_id': rule_id, 'name': name, 'ok': bool(ok), 'detail': detail})


def close(a, b, tol=0.02):
    return a is not None and abs(a - b) < tol


def auto_fix(path, framework):
    """机械可修项：样式 / 页面设置 / updateFields / （--framework）页码分节。返回修复动作列表。"""
    d = docx.Document(path)
    actions = ['setup_styles：Normal/Heading1-4/图题/表题 矫正为规范字号字体行距']
    xju_format.setup_styles(d)
    for i, s in enumerate(d.sections):
        xju_format.setup_section(s)
    actions.append(f'setup_section：{len(d.sections)} 个节的 A4/页边距/页眉页脚距')
    xju_format.enable_update_fields(d)
    actions.append('enable_update_fields：打开时提示更新目录域')
    if framework and len(d.sections) >= 3:
        xju_format.set_pg_num(d.sections[1], 'upperRoman', 1)
        xju_format.set_pg_num(d.sections[2], 'decimal', 1)
        actions.append('set_pg_num：节1 罗马数字 / 节2 阿拉伯数字从 1 起')
    d.save(path)
    return actions


def scan_placeholders(path):
    """扫描正文段落 + 表格单元格里的占位符残留，返回 (count, 前20处示例)。"""
    import re
    d = docx.Document(path)
    pat = re.compile(r'【[^】]*】|\[待填[^\]]*\]|待填|待核实|待定|lorem|ipsum|\bxxx\b', re.I)
    hits = []
    for p in d.paragraphs:
        for m in pat.findall(p.text):
            hits.append({'where': p.text[:40], 'match': m})
    for tbl in d.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for m in pat.findall(cell.text):
                    hits.append({'where': '表格:' + cell.text[:30], 'match': m})
    return len(hits), hits[:20]


def validate(path, framework):
    z = zipfile.ZipFile(path)
    for n in z.namelist():
        if n.endswith('.xml'):
            etree.fromstring(z.read(n))
    check('xml-wellformed', '所有 XML part 良构', True, f'{len(z.namelist())} parts')

    d = docx.Document(path)

    for i, s in enumerate(d.sections):
        check(f'sec{i}-a4', f'节{i} A4 纸张',
              close(s.page_width.cm, 21.0) and close(s.page_height.cm, 29.7))
        check(f'sec{i}-margin', f'节{i} 页边距 上下2.54/左右3.17cm',
              close(s.top_margin.cm, 2.54) and close(s.bottom_margin.cm, 2.54)
              and close(s.left_margin.cm, 3.17) and close(s.right_margin.cm, 3.17),
              f'实际 上{s.top_margin.cm:.2f} 左{s.left_margin.cm:.2f}')
        check(f'sec{i}-hdrftr', f'节{i} 页眉1.5/页脚1.75cm',
              close(s.header_distance.cm, 1.5) and close(s.footer_distance.cm, 1.75))

    def style_sz(name):
        try:
            sz = d.styles[name].font.size
            return sz.pt if sz else None
        except KeyError:
            return None

    for name, want in (('Normal', 12), ('Heading 1', 16), ('Heading 2', 15), ('Heading 3', 14)):
        check(f'style-{name}-size', f'样式 {name} 字号 {want}pt', style_sz(name) == want,
              f'实际 {style_sz(name)}')

    def east_asia(name):
        try:
            rpr = d.styles[name].element.find(qn('w:rPr'))
            rf = rpr.find(qn('w:rFonts')) if rpr is not None else None
            return rf.get(qn('w:eastAsia')) if rf is not None else None
        except KeyError:
            return None

    for name in ('Normal', 'Heading 1', 'Heading 2', 'Heading 3'):
        check(f'style-{name}-font', f'样式 {name} 中文宋体', east_asia(name) == '宋体',
              f'实际 {east_asia(name)}')

    h1 = d.styles['Heading 1']
    check('h1-pagebreak', '一级标题另起一页', bool(h1.paragraph_format.page_break_before))
    check('h1-center', '一级标题居中', str(h1.paragraph_format.alignment) == 'CENTER (1)',
          str(h1.paragraph_format.alignment))

    xml = z.read('word/document.xml').decode('utf-8')
    check('toc-field', '含 TOC 三级目录域', 'TOC \\o' in xml and '1-3' in xml)
    check('update-fields', 'settings 含 updateFields（打开自动提示更新目录）',
          b'updateFields' in z.read('word/settings.xml'))

    hdr_parts = [n for n in z.namelist() if n.startswith('word/header')]
    hdr_ok = any('pBdr' in z.read(n).decode('utf-8') for n in hdr_parts)
    check('header-rule', '页眉存在且带下横线', bool(hdr_parts) and hdr_ok)

    if framework:
        fmts = []
        for s in d.sections:
            pg = s._sectPr.find(qn('w:pgNumType'))
            fmts.append(pg.get(qn('w:fmt')) if pg is not None else None)
        check('pgnum-structure', '三节页码结构 无→罗马→阿拉伯',
              fmts == [None, 'upperRoman', 'decimal'], str(fmts))
        if len(d.sections) >= 3:
            pg = d.sections[2]._sectPr.find(qn('w:pgNumType'))
            check('pgnum-restart', '正文页码从 1 重新起页',
                  pg is not None and pg.get(qn('w:start')) == '1')


MANUAL_CHECKS = ['字体真实渲染效果（Word/WPS 打开看）', '图/表编号分章连续无断号',
                 'GB/T 7714 著录格式细节', '目录更新域后页码正确、引用角标 [1] 保留',
                 '摘要含选题意义/方法/结论、字数达标']


def main():
    args = sys.argv[1:]
    path = args[0]
    framework = '--framework' in args
    as_json = '--json' in args
    do_fix = '--fix' in args
    content_only = '--content' in args

    if content_only:
        count, samples = scan_placeholders(path)
        if as_json:
            print(json.dumps({'status': 'pass' if count == 0 else 'fail',
                              'placeholders': count, 'samples': samples}, ensure_ascii=False))
        else:
            tag = '✅ PASS' if count == 0 else '❌ FAIL'
            print(f'{tag}  占位符残留门禁 — {count} 处')
            for s in samples:
                print(f'    「{s["match"]}」 @ {s["where"]}')
            if count > 20:
                print(f'    …… 仅显示前 20 处，共 {count} 处')
        sys.exit(0 if count == 0 else 1)

    fix_actions = []
    if do_fix:
        fix_actions = auto_fix(path, framework)

    validate(path, framework)
    fails = [r for r in RESULTS if not r['ok']]

    if as_json:
        out = {'status': 'pass' if not fails else 'fail',
               'total_checks': len(RESULTS), 'total_fail': len(fails),
               'fixed': fix_actions,
               'fails': [{'rule_id': r['rule_id'], 'reason': r['name'],
                          'detail': r['detail']} for r in fails[:MAX_JSON_FAILS]],
               'manual_checks': MANUAL_CHECKS}
        if len(fails) > MAX_JSON_FAILS:
            out['truncated'] = f'仅列前 {MAX_JSON_FAILS} 条，先修以上再重试'
        print(json.dumps(out, ensure_ascii=False, indent=1))
    else:
        if fix_actions:
            print('== 已自动修复 ==')
            for a in fix_actions:
                print(f'🔧  {a}')
            print()
        for r in RESULTS:
            tag = '✅ PASS' if r['ok'] else '❌ FAIL'
            print(f'{tag}  {r["name"]}' + (f' — {r["detail"]}' if r['detail'] else ''))
        print()
        if fails:
            print(f'共 {len(fails)} 项不符合规范：{[r["rule_id"] for r in fails]}')
            if not do_fix:
                print('提示：加 --fix 可自动修复样式/页面/updateFields 等机械项。')
        else:
            print('全部机检项通过。人工复核项：' + '；'.join(MANUAL_CHECKS))
    sys.exit(1 if fails else 0)


if __name__ == '__main__':
    main()
