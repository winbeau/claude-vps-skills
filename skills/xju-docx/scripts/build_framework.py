# -*- coding: utf-8 -*-
"""build_framework — 生成符合新疆大学规范的《学年论文》框架 docx（封面+修改记录+目录+九章骨架）。

用法：
  python3 build_framework.py -o 输出.docx [--name 姓名] [--college 学院] \
      [--advisor 指导教师] [--advisor-title 职称] [--title 题目] [--date "2026 年 7 月 9 日"]

章节骨架为《工程项目开发综合实践》课程的 GB8567 文档套件结构；
如需别的骨架，改 STRUCT 或另写生成器复用 xju_format 库。
"""
import argparse
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION

from xju_format import (SIZE, GRAY, set_run_font, spacing_attrs, setup_styles, setup_section,
                        set_pg_num, setup_header, setup_page_footer, front_heading,
                        body_para, placeholder, add_toc)

STRUCT = [
    ('一、引言', [
        (2, '1.1 开发目的', '说明编写本套文档的目的与预期读者（指导教师、团队成员、后续维护者）。'),
        (2, '1.2 现状及意义', '现状痛点 → 系统意义，用平台真实运营数据佐证。'),
        (2, '1.3 背景', '系统名称/线上地址/任务提出者/开发者/用户/运行环境/与外部系统的关系。'),
    ]),
    ('二、系统可行性分析', [
        (2, '2 可行性研究的前提', '一段总起：项目须满足的形态与约束。'),
        (3, '2.1 要求', '功能/性能/输出/输入/处理流程/安全保密/相连系统/完成期限，逐项列出。'),
        (3, '2.2 目标', '信息服务改进、处理速度、人员利用率、控制精度、费用节省。'),
        (3, '2.3 条件、假定和限制', '运行寿命、经费来源、法律政策限制、软硬件环境约束、最晚投用时间。'),
        (3, '2.4 进行可行性研究的方法', '调查分析 + 技术验证 + 方案加权对比 + 敏感性分析。'),
        (3, '2.5 评价尺度', '按优先次序列出评价尺度。'),
        (2, '3 对现有系统的分析', '界定“现有系统”（可为非正式的人工/社交流程）。'),
        (3, '3.1 处理流程和数据流程', '现状流程叙述 + 流程示意图（图 2-1，图题五号宋体加粗居中于图下方）。'),
        (3, '3.2 工作负荷', ''),
        (3, '3.3 费用开支', '直接开支 + 隐性成本量化口径。'),
        (3, '3.4 人员', ''),
        (3, '3.5 设备', ''),
        (3, '3.6 局限性', '逐条列出，并说明改进性维护为何不可行。'),
        (2, '4 所建议的系统', '总起一句。'),
        (3, '4.1 对所建议系统的说明', '功能全景。'),
        (3, '4.2 处理流程和数据流程', '整体数据流 + 数据流图（图 2-2）。'),
        (3, '4.3 改进之处', '与 3.6 局限性逐条对应。'),
        (3, '4.4 影响', '对设备/软件/用户机构/运行过程/开发环境/经费的影响。'),
        (3, '4.5 局限性', '新系统自身局限。'),
        (3, '4.6 技术条件方面的可行性', '技术栈成熟度 + 事实验证。'),
        (2, '5 可选择的其他系统方案', '曾考虑的替代方案及未被推荐的理由。'),
        (3, '5.1 可选择的系统方案 1', ''),
        (3, '5.2 可选择的系统方案 2', ''),
        (2, '6 投资及效益分析', ''),
        (3, '6.1 支出', '列三线表（表 2-1，表题置表上方）。'),
        (3, '6.2 收益', '可货币化 + 不可货币化收益。'),
        (3, '6.3 收益／投资比', ''),
        (3, '6.4 投资回收周期', ''),
        (3, '6.5 敏感性分析', '关键变量变动对结论的影响。'),
        (2, '7 社会因素方面的可行性', ''),
        (3, '7.1 法律方面的可行性', '数据隐私、授权、开源许可义务。'),
        (3, '7.2 使用方面的可行性', ''),
        (2, '8 结论', '可以立即开发/推迟/不可行……的明确结论及依据。'),
    ]),
    ('三、系统开发计划', [
        (2, '1 引言', '编写目的、背景、定义、参考资料。'),
        (2, '2 项目概述', ''),
        (3, '2.1 工作内容', ''),
        (3, '2.2 主要参加人员', ''),
        (3, '2.3 产品', '程序/文档/服务/非移交产品清单。'),
        (3, '2.4 验收标准', ''),
        (3, '2.5 完成项目的最迟期限', ''),
        (3, '2.6 本计划的批准者和批准日期', ''),
        (2, '3 实施计划', ''),
        (3, '3.1 工作任务的分解与人员分工', '按模块分解到人，建议用表格。'),
        (3, '3.2 接口人员', ''),
        (3, '3.3 进度', '甘特图或进度三线表（表 3-1）。'),
        (3, '3.4 预算', '与第二章 6.1 口径一致。'),
        (3, '3.5 关键问题', ''),
        (2, '4 支持条件', ''),
        (3, '4.1 计算机系统支持', ''),
        (3, '4.2 需由用户承担的工作', ''),
        (3, '4.3 由外单位提供的条件', ''),
        (2, '5 专题计划要点', ''),
    ]),
    ('四、系统需求分析', [
        (2, '1．引言', '编写目的、背景、术语定义、参考资料。'),
        (2, '2．任务概述', ''),
        (3, '2.1 目标', ''),
        (2, '3．具体需求分析', ''),
        (3, '3.1 系统流程图', '整体业务流程图（图 4-1）。'),
        (3, '3.2 数据流图及用例模型', '分层 DFD + 按角色的用例图。'),
        (3, '3.3 数据字典', '核心数据项与数据结构定义，用三线表。'),
        (2, '4．支持信息', ''),
        (3, '4.1 运行环境', ''),
        (3, '4.2 支持软件', ''),
        (3, '4.3 接口（外部，内部）', ''),
        (3, '4.4 控制', ''),
        (3, '4.5 需求注释', ''),
        (2, '5．需求分析总结', ''),
    ]),
    ('五、系统概要设计', [
        (2, '1．引言', ''),
        (2, '2．总体设计', ''),
        (3, '2.1 需求规定', ''),
        (3, '2.2 运行环境', ''),
        (3, '2.3 基本设计概念和处理流程', '总体架构图（图 5-1）。'),
        (3, '2.4 软件结构设计', ''),
        (3, '2.5 功能需求与程序的关系', '功能-模块对照三线表。'),
        (3, '2.6 人工处理过程', ''),
        (2, '3．接口设计', ''),
        (3, '3.1 用户接口', ''),
        (3, '3.2 外部接口', ''),
        (3, '3.3 内部接口', ''),
        (2, '4．运行设计', ''),
        (3, '4.1 运行模块组合', ''),
        (3, '4.2 运行控制', ''),
        (2, '5．系统数据库设计', ''),
        (3, '5.1 概念结构设计（E-R 图）', 'E-R 图（图 5-2）。'),
        (3, '5.2 逻辑结构设计', '表结构三线表。'),
        (3, '5.3 物理结构设计', ''),
        (2, '6．系统出错处理设计', ''),
        (3, '6.1 出错信息', ''),
        (3, '6.2 补救措施', '降级与灾备。'),
    ]),
    ('六、系统详细设计', [
        (2, '1．引言', ''),
        (2, '2．程序系统的结构', '程序结构图。'),
        (2, '3．设计说明', ''),
        (3, '3.1 xx 包/类', '按实际模块替换标题。'),
        (4, '3.1.1 xx 函数', '功能/性能/输入输出/算法流程。'),
        (4, '3.1.2 xx 函数', ''),
        (3, '3.2 xx 包/类', ''),
        (2, '4．界面设计', '关键页面截图 + 说明（图 6-x）。'),
        (2, '5. 实现结果', '注出每位同学负责的模块实现结果及核心代码。'),
    ]),
    ('七、系统测试', [
        (2, '1. 引言', ''),
        (2, '2．测试概要', '测试环境、范围、方法。'),
        (2, '3．测试结果及发现', ''),
        (3, '3.1 测试 1（标识符）', '用例-预期-实际，三线表。'),
        (3, '3.2 测试 2（标识符）', ''),
        (2, '4．对软件功能的结论', ''),
        (3, '4.1 功能 1（标识符）', ''),
        (3, '4.2 功能 2（标识符）', ''),
        (2, '5．分析摘要', ''),
        (3, '5.1 能力', ''),
        (3, '5.2 缺陷和限制', ''),
        (3, '5.3 建议', ''),
        (3, '5.4 评价', ''),
        (2, '6．测试资源消耗', ''),
    ]),
    ('八、系统用户操作手册', [
        (2, '1. 引言', ''),
        (3, '1.1 编写目的', ''),
        (3, '1.2 适用范围', ''),
        (3, '1.3 版权声明', ''),
        (2, '2. 安装与卸载', ''),
        (2, '3. 使用说明', ''),
        (3, '3.1 xx 功能', '按功能逐节：截图 + 操作步骤。'),
    ]),
    ('九、总结与展望', [
        (0, '', '归纳成果 → 与既有方式对比 → 学术/应用价值 → 尚存问题与进一步工作。概括、篇幅较短。'),
    ]),
]


def cover_line(doc, label, value, blank_width=10):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(label + '：')
    set_run_font(run, SIZE['三号'])
    pad = max(blank_width - len(value), 2)
    left = pad // 2
    val = '　' * left + value + '　' * (pad - left) if value else '　' * blank_width
    run2 = para.add_run(val)
    set_run_font(run2, SIZE['三号'], underline=True)
    spacing_attrs(para._p.get_or_add_pPr(), line=360)


def build(args):
    doc = Document()
    doc.core_properties.title = '学年论文（工程项目开发综合实践）'
    doc.core_properties.author = args.name or ''
    setup_styles(doc)

    # —— 节1：封面（无页码，页眉从封面开始） ——
    sec1 = doc.sections[0]
    setup_section(sec1)
    setup_header(sec1, args.header)
    sec1.footer.is_linked_to_previous = False   # 显式空页脚

    for _ in range(3):
        body_para(doc, '', indent=False)
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tp.add_run('学　年　论　文')
    set_run_font(run, SIZE['一号'], bold=True)
    sp = doc.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sp.add_run('（工程项目开发综合实践 2025-2026 学年第三学期）')
    set_run_font(run, SIZE['四号'])
    lp = doc.add_paragraph()
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = lp.add_run('（插入团队 Logo）')
    set_run_font(run, SIZE['小四'], color=GRAY)
    body_para(doc, '', indent=False)

    cover_line(doc, '题　　目', args.title or '')
    cover_line(doc, '团　　队', '')
    cover_line(doc, '姓　　名', args.name or '')
    cover_line(doc, '　　　　', '')
    cover_line(doc, '所在学院', args.college, blank_width=12)
    cover_line(doc, '班　　级', '')
    cover_line(doc, '指导教师', args.advisor, blank_width=12)
    cover_line(doc, '职　　称', args.advisor_title, blank_width=12)
    body_para(doc, '', indent=False)
    dp = doc.add_paragraph()
    dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = dp.add_run(args.date)
    set_run_font(run, SIZE['三号'])

    # —— 节2：文档修改记录 + 目录（罗马数字页码） ——
    sec2 = doc.add_section(WD_SECTION.NEW_PAGE)
    setup_section(sec2)
    set_pg_num(sec2, 'upperRoman', 1)
    setup_page_footer(sec2)

    front_heading(doc, ['文档修改记录'], gap='')
    tbl = doc.add_table(rows=11, cols=5)
    tbl.style = 'Table Grid'
    for j, h in enumerate(['序号', '版本号', '更改时间', '更改内容描述', '填写人']):
        cp = tbl.rows[0].cells[j].paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(cp.add_run(h), SIZE['五号'], bold=True)
    for i in range(1, 11):
        vals = [str(i), f'0.{i}' if i < 10 else '1.0', '', '', '']
        for j, v in enumerate(vals):
            cp = tbl.rows[i].cells[j].paragraphs[0]
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_run_font(cp.add_run(v), SIZE['五号'])

    add_toc(doc)

    # —— 节3：正文（阿拉伯数字页码从 1 起，页脚沿用节2 PAGE 域） ——
    sec3 = doc.add_section(WD_SECTION.NEW_PAGE)
    setup_section(sec3)
    set_pg_num(sec3, 'decimal', 1)

    for chap_title, sections in STRUCT:
        doc.add_heading(chap_title, level=1)
        for lvl, title, hint in sections:
            if lvl == 0:
                placeholder(doc, hint)
                continue
            doc.add_heading(title, level=lvl)
            if hint:
                placeholder(doc, hint)

    doc.add_heading('参 考 文 献', level=1)
    for r in ('[1] 主要责任者．文献题名［M］．出版地：出版者，出版年：引文页码．',
              '[2] 析出文献主要责任者．析出文献题名［J］．连续出版物题名，年，卷（期）：页码．',
              '[3] Author A, Author B. Title of the paper[J]. Journal Name, Year, Vol(Issue): Pages.'):
        para = doc.add_paragraph()
        set_run_font(para.add_run(r), SIZE['五号'], color=GRAY)
        spacing_attrs(para._p.get_or_add_pPr(), line=360)
    body_para(doc, '【待填】按 GB/T 7714—2005 著录；文中上角标 [1][2] 实引；不少于 10 篇、外文不少于 2 篇；编号左起顶格，换行与前行文字对齐。', color=GRAY, indent=False)

    doc.add_heading('附录（可选，项目实现部分）团队总结及评分', level=1)
    placeholder(doc, '团队总结及成员评分表；附录中的图表公式另行编号，编号前加“附录1-”字样。')

    doc.save(args.output)
    print('saved:', args.output)


if __name__ == '__main__':
    ap = argparse.ArgumentParser()
    ap.add_argument('-o', '--output', required=True)
    ap.add_argument('--name', default='')
    ap.add_argument('--college', default='计算机科学与技术学院')
    ap.add_argument('--advisor', default='买合木提·买买提')
    ap.add_argument('--advisor-title', default='高级实验师')
    ap.add_argument('--title', default='')
    ap.add_argument('--date', default='2026 年 7 月 9 日')
    ap.add_argument('--header', default='新疆大学学年论文（工程项目开发综合实践）')
    build(ap.parse_args())
