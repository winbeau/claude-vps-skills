# -*- coding: utf-8 -*-
"""xju_format — 按《新疆大学本科毕业论文（设计）规范及格式要求》排版 docx 的 python-docx 工具库。

字号对照（磅）：一号=26 二号=22 三号=16 小三=15 四号=14 小四=12 五号=10.5 小五=9
规范要点：A4；页边距上下 2.54cm、左右 3.17cm；页眉距边 1.5cm、页脚 1.75cm；
一级标题另起一页居中三号宋体单倍行距段前3行段后2行；二级小三顶格；三级四号缩两字符；
正文小四宋体 1.5 倍行距；封面无页码→前置罗马数字→正文阿拉伯数字；
页眉小五宋体居中 + 0.5 磅横线；表格一律三线表（顶/底 1.5 磅、中线 1 磅）。
"""
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SONG, HEI = '宋体', '黑体'
TNR = 'Times New Roman'
GRAY = RGBColor(0x80, 0x80, 0x80)
BLACK = RGBColor(0, 0, 0)

# 中文字号 → 磅
SIZE = {'一号': 26, '二号': 22, '三号': 16, '小三': 15, '四号': 14,
        '小四': 12, '五号': 10.5, '小五': 9, '七号': 5.5}


def set_run_font(run, size, cn=SONG, ascii_=TNR, bold=False, color=BLACK, underline=False):
    run.font.name = ascii_
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.underline = underline
    run._element.rPr.rFonts.set(qn('w:eastAsia'), cn)


def spacing_attrs(pPr, before_lines=None, after_lines=None, before_pt=None, after_pt=None, line=None):
    """w:spacing 直改（beforeLines/afterLines 为 CJK“行”单位×100，python-docx 不暴露）。line: 240=单倍 360=1.5倍"""
    sp = pPr.find(qn('w:spacing'))
    if sp is None:
        sp = OxmlElement('w:spacing')
        pPr.append(sp)
    if before_lines is not None:
        sp.set(qn('w:beforeLines'), str(before_lines))
    if before_pt is not None:
        sp.set(qn('w:before'), str(int(before_pt * 20)))
    if after_lines is not None:
        sp.set(qn('w:afterLines'), str(after_lines))
    if after_pt is not None:
        sp.set(qn('w:after'), str(int(after_pt * 20)))
    if line is not None:
        sp.set(qn('w:line'), str(line))
        sp.set(qn('w:lineRule'), 'auto')


def first_line_chars(pPr, chars=200, fallback_pt=24):
    """首行缩进按“字符”单位（w:firstLineChars），另设磅值兜底。"""
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind')
        pPr.append(ind)
    ind.set(qn('w:firstLineChars'), str(chars))
    ind.set(qn('w:firstLine'), str(int(fallback_pt * 20)))


def style_fonts(st, size, cn=SONG, ascii_=TNR, bold=False):
    st.font.name = ascii_
    st.font.size = Pt(size)
    st.font.bold = bold
    st.font.color.rgb = BLACK
    rpr = st.element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts')
        rpr.insert(0, rf)
    rf.set(qn('w:ascii'), ascii_)
    rf.set(qn('w:hAnsi'), ascii_)
    rf.set(qn('w:eastAsia'), cn)


def _style_outline(st, lvl):
    ppr = st.element.get_or_add_pPr()
    ol = ppr.find(qn('w:outlineLvl'))
    if ol is None:
        ol = OxmlElement('w:outlineLvl')
        ppr.append(ol)
    ol.set(qn('w:val'), str(lvl))


def setup_styles(doc):
    """把 Normal / Heading 1-4 / 图题 / 表题 全部调成规范样式。"""
    normal = doc.styles['Normal']
    style_fonts(normal, SIZE['小四'])
    spacing_attrs(normal.element.get_or_add_pPr(), line=360)

    h1 = doc.styles['Heading 1']          # 另起一页，居中，三号宋体，单倍行距，段前3行段后2行
    style_fonts(h1, SIZE['三号'])
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.page_break_before = True
    spacing_attrs(h1.element.get_or_add_pPr(),
                  before_lines=300, after_lines=200, before_pt=48, after_pt=32, line=240)
    _style_outline(h1, 0)

    h2 = doc.styles['Heading 2']          # 左顶格，小三宋体，1.5倍，段前1行段后0.5行
    style_fonts(h2, SIZE['小三'])
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    spacing_attrs(h2.element.get_or_add_pPr(),
                  before_lines=100, after_lines=50, before_pt=15, after_pt=8, line=360)
    _style_outline(h2, 1)

    h3 = doc.styles['Heading 3']          # 左起空两字符，四号宋体，1.5倍，段前0.5行
    style_fonts(h3, SIZE['四号'])
    h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p = h3.element.get_or_add_pPr()
    spacing_attrs(p, before_lines=50, after_lines=0, before_pt=7, after_pt=0, line=360)
    first_line_chars(p, 200, 28)
    _style_outline(h3, 2)

    h4 = doc.styles['Heading 4']          # 四级（3.1.1 等，不进目录）：小四宋体加粗
    style_fonts(h4, SIZE['小四'], bold=True)
    h4.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p = h4.element.get_or_add_pPr()
    spacing_attrs(p, before_lines=50, after_lines=0, before_pt=6, after_pt=0, line=360)
    first_line_chars(p, 200, 24)
    _style_outline(h4, 3)

    for name in ('图题', '表题'):          # 五号宋体加粗居中，1.5倍，段前段后0行
        try:
            st = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        except ValueError:
            st = doc.styles[name]
        st.base_style = doc.styles['Normal']
        style_fonts(st, SIZE['五号'], bold=True)
        st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        spacing_attrs(st.element.get_or_add_pPr(), before_lines=0, after_lines=0, line=360)


def setup_section(sec):
    """A4 + 规范页边距 + 页眉页脚距边。对文档每个 section 都要调一次。"""
    sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
    sec.top_margin = sec.bottom_margin = Cm(2.54)
    sec.left_margin = sec.right_margin = Cm(3.17)
    sec.header_distance = Cm(1.5)
    sec.footer_distance = Cm(1.75)


def set_pg_num(sec, fmt, start=1):
    """fmt: 'upperRoman'（前置部分）| 'decimal'（正文，start=1 重新起页）。"""
    sectPr = sec._sectPr
    pg = sectPr.find(qn('w:pgNumType'))
    if pg is None:
        pg = OxmlElement('w:pgNumType')
        anchor = sectPr.find(qn('w:cols'))
        if anchor is None:
            anchor = sectPr.find(qn('w:docGrid'))
        if anchor is not None:
            anchor.addprevious(pg)
        else:
            sectPr.append(pg)
    pg.set(qn('w:fmt'), fmt)
    pg.set(qn('w:start'), str(start))


def add_field(para, instr, dirty=False, placeholder_text=None):
    """插入 Word 域（PAGE / TOC …）。TOC 用 dirty=True + settings updateFields 让打开时提示更新。"""
    run = para.add_run()
    fld = OxmlElement('w:fldChar')
    fld.set(qn('w:fldCharType'), 'begin')
    if dirty:
        fld.set(qn('w:dirty'), 'true')
    run._r.append(fld)
    it = OxmlElement('w:instrText')
    it.set(qn('xml:space'), 'preserve')
    it.text = instr
    run._r.append(it)
    sep = OxmlElement('w:fldChar')
    sep.set(qn('w:fldCharType'), 'separate')
    run._r.append(sep)
    if placeholder_text:
        ph = para.add_run(placeholder_text)
        set_run_font(ph, SIZE['小四'], color=GRAY)
    run2 = para.add_run()
    end = OxmlElement('w:fldChar')
    end.set(qn('w:fldCharType'), 'end')
    run2._r.append(end)
    return run


def enable_update_fields(doc):
    settings = doc.settings.element
    if settings.find(qn('w:updateFields')) is None:
        uf = OxmlElement('w:updateFields')
        uf.set(qn('w:val'), 'true')
        settings.append(uf)


def setup_header(sec, text):
    """页眉：小五宋体居中 + 0.5 磅下横线（规范：从封面起每页都有）。"""
    hdr = sec.header
    hdr.is_linked_to_previous = False
    hp = hdr.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = hp.add_run(text)
    set_run_font(run, SIZE['小五'])
    pPr = hp._p.get_or_add_pPr()
    spacing_attrs(pPr, line=240)
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')            # 1/8 磅单位：4 = 0.5 磅
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pbdr.append(bottom)
    pPr.insert(0, pbdr)
    return hdr


def setup_page_footer(sec):
    """页脚：居中 PAGE 域（数字格式由该节 pgNumType 决定）。"""
    ftr = sec.footer
    ftr.is_linked_to_previous = False
    fp = ftr.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_field(fp, ' PAGE ')
    return ftr


def front_heading(doc, chars, gap='　　'):
    """前置部分标题（声明/摘要/目录样式）：三号黑体居中，二字间距两字符，段前3行段后2行。"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(gap.join(chars))
    set_run_font(run, SIZE['三号'], cn=HEI)
    spacing_attrs(para._p.get_or_add_pPr(),
                  before_lines=300, after_lines=200, before_pt=48, after_pt=32, line=240)
    return para


def body_para(doc, text, size=SIZE['小四'], color=BLACK, indent=True, align=None, cn=SONG):
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_run_font(run, size, cn=cn, color=color)
    pPr = para._p.get_or_add_pPr()
    spacing_attrs(pPr, line=360)
    if indent:
        first_line_chars(pPr, 200, size * 2)
    if align is not None:
        para.alignment = align
    return para


def placeholder(doc, text):
    """灰色【待填】占位段。"""
    return body_para(doc, '【待填】' + text, color=GRAY)


def _tbl_border(parent, tag, val, sz):
    el = OxmlElement(tag)
    el.set(qn('w:val'), val)
    if val != 'none':
        el.set(qn('w:sz'), str(sz))
        el.set(qn('w:color'), '000000')
    parent.append(el)


def three_line_table(doc, headers, rows, caption=None):
    """规范三线表：顶/底线 1.5 磅（sz=12）、表头下中线 1 磅（sz=8）、无竖线；
    表题（若给）按“表题”样式置于表上方；表内容五号宋体 1.5 倍行距。"""
    if caption:
        cp = doc.add_paragraph(style='表题')
        run = cp.add_run(caption)
        set_run_font(run, SIZE['五号'], bold=True)
    tbl = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    tblPr = tbl._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    _tbl_border(borders, 'w:top', 'single', 12)
    _tbl_border(borders, 'w:bottom', 'single', 12)
    _tbl_border(borders, 'w:left', 'none', 0)
    _tbl_border(borders, 'w:right', 'none', 0)
    _tbl_border(borders, 'w:insideH', 'none', 0)
    _tbl_border(borders, 'w:insideV', 'none', 0)
    tblPr.append(borders)

    def fill(cell, text, bold=False, header=False):
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        spacing_attrs(p._p.get_or_add_pPr(), before_lines=0, after_lines=0, line=360)
        run = p.add_run(text)
        set_run_font(run, SIZE['五号'], bold=bold)
        if header:                          # 表头下 1 磅中线
            tcPr = cell._tc.get_or_add_tcPr()
            tb = OxmlElement('w:tcBorders')
            _tbl_border(tb, 'w:bottom', 'single', 8)
            tcPr.append(tb)

    for j, h in enumerate(headers):
        fill(tbl.rows[0].cells[j], str(h), bold=True, header=True)
    for i, row in enumerate(rows, start=1):
        for j, v in enumerate(row):
            fill(tbl.rows[i].cells[j], str(v))
    return tbl


def add_toc(doc):
    """目录页：黑体“目 录”标题 + 三级自动目录域（打开文档后更新域生成）。"""
    front_heading(doc, ['目', '录'])
    tocp = doc.add_paragraph()
    add_field(tocp, ' TOC \\o "1-3" \\h \\z \\u ', dirty=True,
              placeholder_text='（目录域：在 Word/WPS 中右键此处→“更新域”，即自动生成三级目录）')
    enable_update_fields(doc)
